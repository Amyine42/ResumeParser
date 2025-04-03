from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

class ResumeFormatter:
    def __init__(self, output_folder):
        self.output_folder = output_folder
        self.styles = getSampleStyleSheet()
        self.setup_styles()
    
    def setup_styles(self):
        # PDF Styles
        if 'Title' not in self.styles:
            self.styles.add(ParagraphStyle(name='Title', fontSize=24, leading=30, alignment=TA_CENTER))
        if 'Subtitle' not in self.styles:
            self.styles.add(ParagraphStyle(name='Subtitle', fontSize=16, leading=24, alignment=TA_CENTER))
        if 'SectionTitle' not in self.styles:
            self.styles.add(ParagraphStyle(name='SectionTitle', fontSize=14, leading=20, spaceBefore=12, spaceAfter=6))
        if 'Normal' not in self.styles:
            self.styles.add(ParagraphStyle(name='Normal', fontSize=12, leading=18))
        if 'Bullet' not in self.styles:
            self.styles.add(ParagraphStyle(name='Bullet', fontSize=12, leading=18, leftIndent=30))
        
        # DOCX Styles
        self.docx_styles = {
            'title': {'size': 24, 'bold': True, 'color': (0, 0, 0)},
            'subtitle': {'size': 16, 'bold': True, 'color': (0, 0, 0)},
            'section_title': {'size': 14, 'bold': True, 'color': (0, 0, 0)},
            'normal': {'size': 12, 'bold': False, 'color': (0, 0, 0)},
            'bullet': {'size': 12, 'bold': False, 'color': (0, 0, 0)}
        }
    
    def generate_pdf(self, data, output_path):
        doc = SimpleDocTemplate(output_path, pagesize=A4)
        elements = []
        
        # Add header
        elements.append(Paragraph(data['header']['name'], self.styles['Title']))
        elements.append(Spacer(1, 12))
        elements.append(Paragraph(data['header']['location'], self.styles['Subtitle']))
        elements.append(Spacer(1, 24))
        
        # Add contact info
        elements.append(Paragraph("Contact Information", self.styles['SectionTitle']))
        contact = []
        for key, value in data['header']['contact'].items():
            if value:
                contact.append(f"{key.title()}: {value}")
        elements.append(Paragraph('<br/>'.join(contact), self.styles['Normal']))
        elements.append(Spacer(1, 24))
        
        # Add experience
        if data['experience']:
            elements.append(Paragraph("Experience", self.styles['SectionTitle']))
            for exp in data['experience']:
                elements.append(Paragraph(exp['title'], self.styles['Normal']))
                elements.append(Paragraph(exp['company'], self.styles['Normal']))
                elements.append(Paragraph(exp['duration'], self.styles['Normal']))
                if exp['description']:
                    elements.append(Paragraph(exp['description'], self.styles['Bullet']))
                elements.append(Spacer(1, 12))
        
        # Add education
        if data['education']:
            elements.append(Paragraph("Education", self.styles['SectionTitle']))
            for edu in data['education']:
                elements.append(Paragraph(edu['degree'], self.styles['Normal']))
                elements.append(Paragraph(edu['institution'], self.styles['Normal']))
                elements.append(Paragraph(edu['year'], self.styles['Normal']))
                elements.append(Spacer(1, 12))
        
        # Add skills
        if data['skills']:
            elements.append(Paragraph("Skills", self.styles['SectionTitle']))
            elements.append(Paragraph(', '.join(data['skills']), self.styles['Normal']))
            elements.append(Spacer(1, 12))
        
        # Add certifications
        if data['certifications']:
            elements.append(Paragraph("Certifications", self.styles['SectionTitle']))
            for cert in data['certifications']:
                elements.append(Paragraph(cert, self.styles['Normal']))
                elements.append(Spacer(1, 12))
        
        doc.build(elements)
    
    def generate_docx(self, data, output_path):
        doc = Document()
        
        # Add header
        title = doc.add_heading(level=1)
        title.add_run(data['header']['name']).font.size = Pt(self.docx_styles['title']['size'])
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading(level=2)
        subtitle.add_run(data['header']['location']).font.size = Pt(self.docx_styles['subtitle']['size'])
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Add contact info
        doc.add_heading("Contact Information", level=2)
        contact = doc.add_paragraph()
        for key, value in data['header']['contact'].items():
            if value:
                contact.add_run(f"{key.title()}: {value}\n")
        
        # Add experience
        if data['experience']:
            doc.add_heading("Experience", level=2)
            for exp in data['experience']:
                doc.add_paragraph(exp['title'], style='Heading 3')
                doc.add_paragraph(exp['company'])
                doc.add_paragraph(exp['duration'])
                if exp['description']:
                    doc.add_paragraph(exp['description'], style='List Bullet')
        
        # Add education
        if data['education']:
            doc.add_heading("Education", level=2)
            for edu in data['education']:
                doc.add_paragraph(edu['degree'], style='Heading 3')
                doc.add_paragraph(edu['institution'])
                doc.add_paragraph(edu['year'])
        
        # Add skills
        if data['skills']:
            doc.add_heading("Skills", level=2)
            doc.add_paragraph(', '.join(data['skills']))
        
        # Add certifications
        if data['certifications']:
            doc.add_heading("Certifications", level=2)
            for cert in data['certifications']:
                doc.add_paragraph(cert, style='List Bullet')
        
        doc.save(output_path)
    
    def format_resume(self, data, output_format='pdf'):
        if output_format == 'pdf':
            output_path = os.path.join(self.output_folder, f"formatted_{data['filename']}.pdf")
            self.generate_pdf(data, output_path)
        else:
            output_path = os.path.join(self.output_folder, f"formatted_{data['filename']}.docx")
            self.generate_docx(data, output_path)
        
        return output_path