from docx import Document
import os
from config import Config

class DocxHandler:
    def __init__(self, file_path):
        self.file_path = file_path
        self.text = ""
    
    def extract_text(self):
        doc = Document(self.file_path)
        for paragraph in doc.paragraphs:
            self.text += paragraph.text + "\n"
        return self.text
    
    def extract_tables(self):
        doc = Document(self.file_path)
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                table_data.append(row_data)
            tables.append(table_data)
        return tables